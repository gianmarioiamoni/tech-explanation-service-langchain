# ui/utils/document_exporter.py
# Service for exporting chat content to various document formats

import os
import tempfile
from datetime import datetime
from typing import Tuple

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentExporter:
    # Export chat content to various document formats (MD, PDF, DOCX)
    #
    # Args:
    #     None
    #
    # Returns:
    #     None
    
    @staticmethod
    def export_to_markdown(topic: str, content: str) -> Tuple[str, str]:
        # Export content to Markdown format
        #
        # Args:
        #     topic: Chat topic/title
        #     content: Explanation text
        #
        # Returns:
        #     Tuple of (file_path, filename)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"tech_explanation_{timestamp}.md"
        
        # Create markdown content
        md_content = f"# {topic}\n\n"
        md_content += f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n"
        md_content += "---\n\n"
        md_content += content
        
        # Write to temp file
        temp_path = os.path.join(tempfile.gettempdir(), filename)
        with open(temp_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return temp_path, filename
    
    @staticmethod
    def export_to_pdf(topic: str, content: str) -> Tuple[str, str]:
        # Export content to PDF format using ReportLab (full Unicode support)
        #
        # Args:
        #     topic: Chat topic/title
        #     content: Explanation text
        #
        # Returns:
        #     Tuple of (file_path, filename)
        
        if not PDF_AVAILABLE:
            raise ImportError("reportlab not installed. Install with: pip install reportlab")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"tech_explanation_{timestamp}.pdf"
        temp_path = os.path.join(tempfile.gettempdir(), filename)
        
        # Create PDF document
        doc = SimpleDocTemplate(
            temp_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )
        
        # Container for PDF elements
        story = []
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='black',
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold',
        )
        
        timestamp_style = ParagraphStyle(
            'Timestamp',
            parent=styles['Normal'],
            fontSize=10,
            textColor='grey',
            alignment=TA_CENTER,
            fontName='Helvetica-Oblique',
        )
        
        body_style = styles['BodyText']
        body_style.fontSize = 11
        body_style.leading = 14
        
        # Add title
        story.append(Paragraph(topic, title_style))
        
        # Add timestamp
        timestamp_text = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        story.append(Paragraph(timestamp_text, timestamp_style))
        story.append(Spacer(1, 0.3 * inch))
        
        # Add content paragraphs
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                # Replace single newlines with <br/> tags for line breaks
                formatted_text = paragraph.replace('\n', '<br/>')
                story.append(Paragraph(formatted_text, body_style))
                story.append(Spacer(1, 0.1 * inch))
        
        # Build PDF
        doc.build(story)
        
        return temp_path, filename
    
    @staticmethod
    def export_to_docx(topic: str, content: str) -> Tuple[str, str]:
        # Export content to Word DOCX format
        #
        # Args:
        #     topic: Chat topic/title
        #     content: Explanation text
        #
        # Returns:
        #     Tuple of (file_path, filename)
        
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"tech_explanation_{timestamp}.docx"
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading(topic, level=1)
        title.alignment = 1  # Center alignment
        
        # Timestamp
        timestamp_para = doc.add_paragraph()
        timestamp_run = timestamp_para.add_run(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        timestamp_run.italic = True
        timestamp_run.font.size = Pt(10)
        timestamp_para.alignment = 1  # Center alignment
        
        # Add spacing
        doc.add_paragraph()
        
        # Content
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        # Save to temp file
        temp_path = os.path.join(tempfile.gettempdir(), filename)
        doc.save(temp_path)
        
        return temp_path, filename
    
    @staticmethod
    def export_chat(topic: str, content: str, format: str) -> Tuple[str, str]:
        # Export chat to specified format
        #
        # Args:
        #     topic: Chat topic/title
        #     content: Explanation text
        #     format: Export format ("Markdown", "PDF", "Word")
        #
        # Returns:
        #     Tuple of (file_path, filename)
        #
        # Raises:
        #     ValueError: If format is not supported
        #     ImportError: If required library is not installed
        
        if not content or not content.strip():
            raise ValueError("No content to export")
        
        if format == "Markdown":
            return DocumentExporter.export_to_markdown(topic, content)
        elif format == "PDF":
            return DocumentExporter.export_to_pdf(topic, content)
        elif format == "Word":
            return DocumentExporter.export_to_docx(topic, content)
        else:
            raise ValueError(f"Unsupported format: {format}")

